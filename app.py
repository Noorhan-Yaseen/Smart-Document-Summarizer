from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
import os
import time
import logging
from werkzeug.utils import secure_filename
from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document
from docx.shared import Pt
from transformers import BartForConditionalGeneration, BartTokenizer
from io import BytesIO

# Initialize Flask
app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx'}
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load BART model
try:
    logger.info("Loading BART model...")
    model = BartForConditionalGeneration.from_pretrained('facebook/bart-large-cnn')
    tokenizer = BartTokenizer.from_pretrained('facebook/bart-large-cnn')
    logger.info("Model loaded successfully!")
except Exception as e:
    logger.error(f"Model loading failed: {e}")
    raise

# Utility Functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def extract_text_from_pdf(filepath):
    try:
        return extract_pdf_text(filepath)
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""

def extract_text_from_docx(filepath):
    try:
        doc = Document(filepath)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""

def generate_summary(text):
    try:
        inputs = tokenizer([text], max_length=1024, truncation=True, return_tensors="pt")
        summary_ids = model.generate(
            inputs['input_ids'],
            num_beams=4,
            max_length=200,
            min_length=50,
            early_stopping=True
        )
        return tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    except Exception as e:
        logger.error(f"Summarization failed: {e}")
        return "Summary generation failed."

def create_docx(summary_text):
    doc = Document()
    doc.add_heading('Document Summary', 0)
    paragraph = doc.add_paragraph(summary_text)
    paragraph.style.font.size = Pt(12)
    return doc

# Routes
@app.route('/')
def home():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    files = request.files.getlist('file')

    if not files or all(file.filename.strip() == '' for file in files):
        return jsonify({"error": "No selected files"}), 400

    results = []

    for file in files:
        filename = secure_filename(file.filename)

        if not allowed_file(filename):
            results.append({
                "filename": filename,
                "error": "File type not allowed"
            })
            continue

        try:
            # Save uploaded file
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = extract_text_from_pdf(filepath)
            elif filename.lower().endswith('.docx'):
                text = extract_text_from_docx(filepath)
            else:
                text = ""

            if not text.strip():
                results.append({
                    "filename": filename,
                    "error": "No text found in document"
                })
                continue

            # Generate summary
            summary = generate_summary(text)

            # Create unique filenames
            timestamp = str(int(time.time()))
            base_name = f"summary_{filename.rsplit('.', 1)[0]}_{timestamp}"

            # Save summary to TXT
            txt_filename = f"{base_name}.txt"
            txt_path = os.path.join(app.config['OUTPUT_FOLDER'], txt_filename)
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(summary)

            # Save summary to DOCX
            docx_filename = f"{base_name}.docx"
            docx_path = os.path.join(app.config['OUTPUT_FOLDER'], docx_filename)
            doc = create_docx(summary)
            doc.save(docx_path)

            # Remove uploaded file after processing
            os.remove(filepath)

            # Append success result
            results.append({
                "filename": filename,
                "summary": summary,
                "original_length": len(text),
                "summary_length": len(summary),
                "download_txt": f"/download/{txt_filename}",
                "download_docx": f"/download/{docx_filename}"
            })

        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            results.append({
                "filename": filename,
                "error": str(e)
            })

    # Final response
    if all('error' in result for result in results):
        return jsonify({"error": "All files failed to process", "details": results}), 400

    return jsonify({"results": results})

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    try:
        filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)

        if not os.path.exists(filepath):
            return jsonify({"error": "File not found"}), 404

        if filename.endswith('.txt'):
            mimetype = 'text/plain'
        elif filename.endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'application/octet-stream'

        return send_file(filepath, mimetype=mimetype, as_attachment=True, download_name=filename)

    except Exception as e:
        logger.error(f"Error downloading file: {e}")
        return jsonify({"error": str(e)}), 500

# Run the app
if __name__ == '__main__':
    app.run(debug=True)